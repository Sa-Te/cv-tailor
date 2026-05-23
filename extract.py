"""
Extract plain text from an uploaded CV file (PDF or DOCX).

Streamlit hands us an UploadedFile (a file-like object), so everything here
reads from a stream rather than a path.
"""

from io import BytesIO
from pypdf import PdfReader
from docx import Document


def extract_text(uploaded_file) -> str:
    """
    Take a Streamlit UploadedFile and return its text content.
    Supports .pdf and .docx. Raises ValueError for anything else.
    """
    name = uploaded_file.name.lower()
    data = uploaded_file.read()

    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    raise ValueError(
        f"Unsupported file type: {uploaded_file.name}. Please upload a PDF or DOCX."
    )


def _from_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts).strip()


def _from_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # also pull text out of any tables (some CVs use them for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()
