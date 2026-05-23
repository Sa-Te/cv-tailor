"""
Render a tailored-CV dict into three formats: Markdown, DOCX, PDF.

All three are deliberately single-column, no tables, no graphics — that's
what keeps them ATS-parseable. Pretty multi-column CVs look nice to humans
and confuse the parsers, so we don't do them here.
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import mm


def _contact_line(contact: dict) -> str:
    parts = [
        contact.get("location"),
        contact.get("phone"),
        contact.get("email"),
        contact.get("linkedin"),
        contact.get("portfolio"),
    ]
    return " | ".join(p for p in parts if p)


# --------------------------------------------------------------------------- #
# Markdown
# --------------------------------------------------------------------------- #
def to_markdown(cv: dict) -> str:
    md = [f"# {cv.get('name', '')}\n\n"]
    md.append(_contact_line(cv.get("contact", {})) + "\n")

    if cv.get("summary"):
        md.append("\n## Professional Summary\n\n")
        md.append(cv["summary"] + "\n")

    if cv.get("skills"):
        md.append("\n## Key Skills\n\n")
        md.append(", ".join(cv["skills"]) + "\n")

    if cv.get("experience"):
        md.append("\n## Professional Experience\n")
        for job in cv["experience"]:
            md.append(
                f"\n**{job.get('title','')}** — {job.get('company','')}, "
                f"{job.get('location','')}  \n"
            )
            md.append(f"*{job.get('start_date','')} – {job.get('end_date','')}*\n\n")
            for b in job.get("bullets", []):
                md.append(f"- {b}\n")

    if cv.get("education"):
        md.append("\n## Education\n\n")
        for e in cv["education"]:
            line = f"- **{e.get('qualification','')}** — {e.get('institution','')} ({e.get('year','')})"
            if e.get("details"):
                line += f" — {e['details']}"
            md.append(line + "\n")

    if cv.get("certifications"):
        md.append("\n## Certifications\n\n")
        for cert in cv["certifications"]:
            md.append(f"- {cert}\n")

    return "".join(md)


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def to_docx(cv: dict) -> BytesIO:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(16)

    # Contact
    contact_p = doc.add_paragraph(_contact_line(cv.get("contact", {})))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def section(title):
        p = doc.add_paragraph()
        r = p.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(12)

    if cv.get("summary"):
        section("Professional Summary")
        doc.add_paragraph(cv["summary"])

    if cv.get("skills"):
        section("Key Skills")
        doc.add_paragraph(" • ".join(cv["skills"]))

    if cv.get("experience"):
        section("Professional Experience")
        for job in cv["experience"]:
            head = doc.add_paragraph()
            head.add_run(
                f"{job.get('title','')} — {job.get('company','')}, {job.get('location','')}"
            ).bold = True
            date_p = doc.add_paragraph()
            date_run = date_p.add_run(f"{job.get('start_date','')} – {job.get('end_date','')}")
            date_run.italic = True
            for b in job.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    if cv.get("education"):
        section("Education")
        for e in cv["education"]:
            line = f"{e.get('qualification','')} — {e.get('institution','')} ({e.get('year','')})"
            if e.get("details"):
                line += f" — {e['details']}"
            doc.add_paragraph(line)

    if cv.get("certifications"):
        section("Certifications")
        for cert in cv["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------- #
# PDF (ReportLab — cross-platform, no Word/LibreOffice dependency)
# --------------------------------------------------------------------------- #
def _esc(text: str) -> str:
    """Escape characters that ReportLab's mini-HTML would choke on."""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def to_pdf(cv: dict) -> BytesIO:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=16, alignment=TA_CENTER, spaceAfter=2)
    centre = ParagraphStyle("centre", parent=styles["BodyText"], alignment=TA_CENTER, spaceAfter=8)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("body", parent=styles["BodyText"], spaceAfter=2)

    story = [Paragraph(_esc(cv.get("name", "")), h1)]
    story.append(Paragraph(_esc(_contact_line(cv.get("contact", {}))), centre))

    if cv.get("summary"):
        story.append(Paragraph("PROFESSIONAL SUMMARY", h2))
        story.append(Paragraph(_esc(cv["summary"]), body))

    if cv.get("skills"):
        story.append(Paragraph("KEY SKILLS", h2))
        story.append(Paragraph(_esc(" • ".join(cv["skills"])), body))

    if cv.get("experience"):
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", h2))
        for job in cv["experience"]:
            story.append(
                Paragraph(
                    f"<b>{_esc(job.get('title',''))}</b> — "
                    f"{_esc(job.get('company',''))}, {_esc(job.get('location',''))}",
                    body,
                )
            )
            story.append(
                Paragraph(f"<i>{_esc(job.get('start_date',''))} – {_esc(job.get('end_date',''))}</i>", body)
            )
            for b in job.get("bullets", []):
                story.append(Paragraph(f"• {_esc(b)}", body))
            story.append(Spacer(1, 6))

    if cv.get("education"):
        story.append(Paragraph("EDUCATION", h2))
        for e in cv["education"]:
            line = f"<b>{_esc(e.get('qualification',''))}</b> — {_esc(e.get('institution',''))} ({_esc(e.get('year',''))})"
            if e.get("details"):
                line += f" — {_esc(e['details'])}"
            story.append(Paragraph(line, body))

    if cv.get("certifications"):
        story.append(Paragraph("CERTIFICATIONS", h2))
        for cert in cv["certifications"]:
            story.append(Paragraph(f"• {_esc(cert)}", body))

    doc.build(story)
    buf.seek(0)
    return buf
